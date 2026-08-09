"""대본을 .docx로 만들어 바이트로 돌려준다.

피그마가 두 군데서 요구한다.
- AI Script Edit Page ㉒: "13의 제목명.docx로 저장" (편집한 대본 그대로)
- Coach View Page ㉙: "하이라이팅_대본.docx로 저장" (발음 주의 단어에 색이 칠해진 대본)

## 왜 서버에서 만드나
`python-docx`가 이미 의존성에 있다(지금까지는 업로드된 docx를 *읽는* 데만 썼다). 프론트에서
docx를 만들려면 라이브러리를 새로 넣어야 하고 한글 폰트 처리가 붙는다. 무엇보다 하이라이팅
색상은 서버가 분류한 카테고리에서 나오므로, 색을 칠할 근거가 서버에 있다.
"""
import io

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

# 피그마 Coach View Page ㉜에 적힌 값 그대로.
#   장단음 F7358E / 연음 0072F2 / 표기-발음 불일치 F79322
CATEGORY_COLORS = {
    "장단음": RGBColor(0xF7, 0x35, 0x8E),
    "연음": RGBColor(0x00, 0x72, 0xF2),
    "표기-발음불일치": RGBColor(0xF7, 0x93, 0x22),
}

# 한글이 깨지지 않는 기본 폰트. 워드가 없으면 시스템 기본으로 대체된다.
_BODY_FONT = "맑은 고딕"


def _new_document(title: str):
    document = docx.Document()
    style = document.styles["Normal"]
    style.font.name = _BODY_FONT
    style.font.size = Pt(11)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    return document


def _to_bytes(document) -> bytes:
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_script_docx(title: str, slides) -> bytes:
    """편집한 대본을 그대로 담는다.

    slides: [(slide_number, script)] 순서대로. 슬라이드가 하나뿐이면(PPT 없이 만든 프로젝트)
    "Slide 1" 머리글을 붙이지 않는다 — 한 덩어리 대본에 번호를 붙이면 어색하다.
    """
    document = _new_document(title)
    single = len(slides) <= 1

    for slide_number, script in slides:
        if not single:
            head = document.add_paragraph()
            run = head.add_run(f"Slide {slide_number}")
            run.bold = True
            run.font.size = Pt(12)
        for line in (script or "").split("\n"):
            document.add_paragraph(line)
        if not single:
            document.add_paragraph()

    return _to_bytes(document)


def build_highlighted_script_docx(title: str, slides, difficult_words) -> bytes:
    """발음 주의 단어에 카테고리 색을 칠한 대본.

    difficult_words: [{"word", "phoneme", "category"}]. 같은 단어가 여러 번 나오면 전부 칠한다
    (피그마 ㉛도 "같은 단어가 여러 번 나올 경우 <>로 찾기 가능"이라 모두 표시가 전제다).

    긴 단어를 먼저 찾는다 — '발표'와 '발표자'가 둘 다 목록에 있으면 짧은 쪽이 먼저 걸려서
    긴 단어가 쪼개진다.
    """
    document = _new_document(title)

    lookup = {}
    for item in difficult_words or []:
        word = (item.get("word") or "").strip()
        if word and item.get("category") in CATEGORY_COLORS:
            lookup[word] = item
    targets = sorted(lookup, key=len, reverse=True)

    single = len(slides) <= 1
    for slide_number, script in slides:
        if not single:
            head = document.add_paragraph()
            run = head.add_run(f"Slide {slide_number}")
            run.bold = True
            run.font.size = Pt(12)
        for line in (script or "").split("\n"):
            _write_highlighted_line(document, line, targets, lookup)
        if not single:
            document.add_paragraph()

    _append_legend(document, difficult_words, lookup)
    return _to_bytes(document)


def _write_highlighted_line(document, line: str, targets, lookup) -> None:
    paragraph = document.add_paragraph()
    if not line or not targets:
        paragraph.add_run(line)
        return

    index = 0
    while index < len(line):
        match = None
        for word in targets:
            if line.startswith(word, index):
                match = word
                break
        if match:
            run = paragraph.add_run(match)
            run.font.color.rgb = CATEGORY_COLORS[lookup[match]["category"]]
            run.bold = True
            index += len(match)
        else:
            # 한 글자씩 붙이면 run이 폭발하므로, 다음 후보 시작 위치까지 한 번에 붙인다.
            next_hit = len(line)
            for word in targets:
                found = line.find(word, index + 1)
                if found != -1:
                    next_hit = min(next_hit, found)
            paragraph.add_run(line[index:next_hit])
            index = next_hit


def _append_legend(document, difficult_words, lookup) -> None:
    """어떤 색이 무슨 뜻인지 없으면 색칠이 무의미하다. 실제로 칠해진 단어만 적는다."""
    if not lookup:
        return

    document.add_paragraph()
    head = document.add_paragraph()
    run = head.add_run("발음 주의 단어")
    run.bold = True
    run.font.size = Pt(12)

    for item in difficult_words or []:
        word = (item.get("word") or "").strip()
        if word not in lookup:
            continue
        paragraph = document.add_paragraph()
        colored = paragraph.add_run(f"{word} ")
        colored.font.color.rgb = CATEGORY_COLORS[item["category"]]
        colored.bold = True
        paragraph.add_run(f"{item.get('phoneme') or ''} · {item['category']}")
        description = (item.get("description") or "").strip()
        if description:
            paragraph.add_run(f" — {description}")
